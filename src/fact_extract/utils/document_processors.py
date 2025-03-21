"""
Document processors for different file types.
Handles extraction of text content from various document formats.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging
import pandas as pd
from openpyxl import load_workbook
from docx import Document
from pypdf import PdfReader
import re
import csv

logger = logging.getLogger(__name__)

class DocumentProcessor(ABC):
    """Base class for document processors."""
    
    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if this processor can handle the given file.
        
        Args:
            file_path: Path to the document
            
        Returns:
            bool: True if this processor can handle the file
        """
        pass
        
    @abstractmethod
    def extract_content(self, file_path: Path) -> List[Dict[str, str]]:
        """Extract content from the document.
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of dictionaries containing:
                - title: Document or section title
                - content: Text content
                - source: Original file path
        """
        pass
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text by removing extra whitespace etc.
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?-]', '', text)
        return text.strip()

class ExcelProcessor(DocumentProcessor):
    """Process Excel files."""
    
    def __init__(self, title_columns: List[str] = None, content_columns: List[str] = None):
        """Initialize Excel processor.
        
        Args:
            title_columns: List of possible column names for titles
            content_columns: List of possible column names for content
        """
        self.title_columns = title_columns or ['title', 'name', 'heading', 'subject']
        self.content_columns = content_columns or ['content', 'text', 'body', 'description']
        
    def can_process(self, file_path: Path) -> bool:
        """Check if file is Excel format."""
        return file_path.suffix.lower() in ['.xlsx', '.xls']
        
    def extract_content(self, file_path: Path) -> List[Dict[str, str]]:
        """Extract content from Excel file."""
        try:
            # Read Excel file
            df = pd.read_excel(file_path)
            
            # Find title and content columns
            title_col = next((col for col in df.columns if col.lower() in self.title_columns), None)
            content_col = next((col for col in df.columns if col.lower() in self.content_columns), None)
            
            if not content_col:
                logger.warning(f"No content column found in {file_path}")
                return []
                
            results = []
            for _, row in df.iterrows():
                content = str(row[content_col])
                if pd.isna(content) or not content.strip():
                    continue
                    
                title = str(row[title_col]) if title_col else "Untitled"
                if pd.isna(title):
                    title = "Untitled"
                    
                results.append({
                    'title': self._clean_text(title),
                    'content': self._clean_text(content),
                    'source': str(file_path)
                })
                
            return results
            
        except Exception as e:
            logger.error(f"Error processing Excel file {file_path}: {str(e)}")
            return []

class WordProcessor(DocumentProcessor):
    """Process Word documents."""
    
    def can_process(self, file_path: Path) -> bool:
        """Check if file is Word format."""
        return file_path.suffix.lower() in ['.docx', '.doc']
        
    def extract_content(self, file_path: Path) -> List[Dict[str, str]]:
        """Extract content from Word file."""
        try:
            doc = Document(file_path)
            
            # Extract title from first paragraph if it looks like a title
            title = "Untitled"
            content_start = 0
            
            # Look for title in the first few paragraphs
            for i, para in enumerate(doc.paragraphs[:3]):
                # Check if paragraph style indicates a title
                style_name = para.style.name.lower()
                if ('heading' in style_name and '1' in style_name) or 'title' in style_name:
                    title = para.text
                    content_start = i + 1
                    break
                # Check if paragraph looks like a title (short, all caps, etc.)
                elif (len(para.text.split()) <= 10 and 
                      para.text.strip() and 
                      not para.text.endswith('.')):
                    title = para.text
                    content_start = i + 1
                    break
            
            # Combine remaining paragraphs into content
            content = []
            for para in doc.paragraphs[content_start:]:
                if para.text.strip():
                    content.append(para.text.strip())
            
            if not content:
                return []
                
            return [{
                'title': self._clean_text(title),
                'content': self._clean_text('\n'.join(content)),
                'source': str(file_path)
            }]
            
        except Exception as e:
            logger.error(f"Error processing Word file {file_path}: {str(e)}")
            return []

class PDFProcessor(DocumentProcessor):
    """Process PDF documents."""
    
    def can_process(self, file_path: Path) -> bool:
        """Check if file is PDF format."""
        return file_path.suffix.lower() == '.pdf'
        
    def extract_content(self, file_path: Path) -> List[Dict[str, str]]:
        """Extract content from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                
                # Try to get title from metadata
                title = "Untitled"
                if reader.metadata:
                    if '/Title' in reader.metadata:
                        title = reader.metadata['/Title']
                    elif 'title' in reader.metadata:
                        title = reader.metadata['title']
                
                # Extract text from all pages
                content = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        # Clean up PDF-specific artifacts
                        text = re.sub(r'\s*\n\s*', ' ', text)
                        text = re.sub(r'\s+', ' ', text)
                        content.append(text.strip())
                
                if not content:
                    return []
                    
                return [{
                    'title': self._clean_text(title),
                    'content': self._clean_text('\n'.join(content)),
                    'source': str(file_path)
                }]
                
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            return []

class CSVProcessor(DocumentProcessor):
    """Process CSV files."""
    
    def __init__(self, title_columns: List[str] = None, content_columns: List[str] = None):
        """Initialize CSV processor.
        
        Args:
            title_columns: List of possible column names for titles
            content_columns: List of possible column names for content
        """
        self.title_columns = title_columns or ['title', 'name', 'heading', 'subject']
        self.content_columns = content_columns or ['content', 'text', 'body', 'description']
        
    def can_process(self, file_path: Path) -> bool:
        """Check if file is CSV format."""
        return file_path.suffix.lower() == '.csv'
        
    def extract_content(self, file_path: Path) -> List[Dict[str, str]]:
        """Extract content from CSV file.
        
        This processor handles:
        1. Auto-detection of CSV dialect
        2. Multiple character encodings (utf-8, latin-1, etc.)
        3. Flexible column name matching
        4. Row-by-row processing for memory efficiency
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    # Try to read with pandas directly first
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except (UnicodeDecodeError, pd.errors.ParserError):
                    try:
                        # If pandas fails, try manual dialect detection
                        with open(file_path, 'r', encoding=encoding) as f:
                            sample = f.read(1024)
                            dialect = csv.Sniffer().sniff(sample)
                            has_header = csv.Sniffer().has_header(sample)
                            
                            # Rewind and read with detected dialect
                            f.seek(0)
                            reader = csv.reader(f, dialect)
                            
                            # Convert to DataFrame
                            data = list(reader)
                            if has_header:
                                columns = data[0]
                                data = data[1:]
                            else:
                                columns = [f'Column{i}' for i in range(len(data[0]))]
                                
                            df = pd.DataFrame(data, columns=columns)
                            break
                    except Exception:
                        continue
            
            if df is None:
                logger.error(f"Could not read CSV file {file_path} with any supported encoding")
                return []
            
            # Find title and content columns (case-insensitive)
            df.columns = [str(col).lower() for col in df.columns]
            title_col = next((col for col in df.columns if col in self.title_columns), None)
            content_col = next((col for col in df.columns if col in self.content_columns), None)
            
            if not content_col:
                logger.warning(f"No content column found in {file_path}")
                return []
                
            results = []
            for _, row in df.iterrows():
                content = str(row[content_col])
                if pd.isna(content) or not content.strip():
                    continue
                    
                title = str(row[title_col]) if title_col else "Untitled"
                if pd.isna(title):
                    title = "Untitled"
                    
                results.append({
                    'title': self._clean_text(title),
                    'content': self._clean_text(content),
                    'source': str(file_path)
                })
                
            return results
            
        except Exception as e:
            logger.error(f"Error processing CSV file {file_path}: {str(e)}")
            return []

class DocumentProcessorFactory:
    """Factory for creating document processors."""
    
    def __init__(self):
        """Initialize with default processors."""
        self.processors = [
            ExcelProcessor(),
            CSVProcessor(),
            WordProcessor(),
            PDFProcessor()
        ]
    
    def get_processor(self, file_path: Path) -> Optional[DocumentProcessor]:
        """Get appropriate processor for file type.
        
        Args:
            file_path: Path to document
            
        Returns:
            DocumentProcessor if supported, None otherwise
        """
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        return None 