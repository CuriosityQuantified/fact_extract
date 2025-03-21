"""
Edge case tests for document processors.
Tests various edge cases and error conditions.
"""

import os
import sys
import pytest
import pandas as pd
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from pypdf import PdfReader
from reportlab.pdfgen import canvas
from io import BytesIO
import csv


# Ensure the src directory is in the path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Ensure the src directory is in the path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '.')))
from utils.document_processors import (
    DocumentProcessorFactory,
    ExcelProcessor,
    WordProcessor,
    PDFProcessor,
    CSVProcessor
)

# Test data directory
TEST_DATA_DIR = Path("test_data/edge_cases")

def setup_module(module):
    """Set up test data before running tests."""
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    
    # Create edge case test files
    create_empty_files()
    create_malformed_files()
    create_large_files()
    create_special_character_files()
    create_mixed_encoding_files()

def create_empty_files():
    """Create empty files for testing."""
    # Empty files
    (TEST_DATA_DIR / 'empty.xlsx').touch()
    (TEST_DATA_DIR / 'empty.csv').touch()
    (TEST_DATA_DIR / 'empty.docx').write_bytes(b'PK\x03\x04\x14\x00\x00\x00\x00\x00')  # Minimal DOCX
    
    # Create empty PDF
    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    c.save()
    with open(TEST_DATA_DIR / 'empty.pdf', 'wb') as f:
        f.write(buffer.getvalue())

def create_malformed_files():
    """Create malformed files for testing."""
    # Malformed Excel (invalid XML)
    with open(TEST_DATA_DIR / 'malformed.xlsx', 'wb') as f:
        f.write(b'Invalid Excel Content')
    
    # Malformed CSV (mixed delimiters)
    with open(TEST_DATA_DIR / 'malformed.csv', 'w') as f:
        f.write('col1,col2;col3|col4\nval1;val2,val3|val4')
    
    # Malformed Word (corrupted content)
    with open(TEST_DATA_DIR / 'malformed.docx', 'wb') as f:
        f.write(b'Invalid Word Content')
    
    # Malformed PDF (corrupted structure)
    with open(TEST_DATA_DIR / 'malformed.pdf', 'wb') as f:
        f.write(b'%PDF-1.7\nInvalid PDF Content')

def create_large_files():
    """Create large files for testing memory handling."""
    # Large Excel with many rows
    df = pd.DataFrame({
        'title': [f'Title {i}' for i in range(10000)],
        'content': [f'Content {i} ' * 100 for i in range(10000)]
    })
    df.to_excel(TEST_DATA_DIR / 'large.xlsx', index=False)
    
    # Large CSV with many rows
    df.to_csv(TEST_DATA_DIR / 'large.csv', index=False)
    
    # Large Word document with many paragraphs
    doc = Document()
    doc.add_heading('Large Document', 0)
    for i in range(1000):
        doc.add_paragraph(f'Paragraph {i} ' * 50)
    doc.save(TEST_DATA_DIR / 'large.docx')
    
    # Large PDF with many pages
    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    for i in range(100):
        c.drawString(100, 750, f'Page {i} content ' * 20)
        c.showPage()
    c.save()
    with open(TEST_DATA_DIR / 'large.pdf', 'wb') as f:
        f.write(buffer.getvalue())

def create_special_character_files():
    """Create files with special characters."""
    # Excel with special characters
    df = pd.DataFrame({
        'title': ['Title with 特殊字符', 'Title with ñáéíóú', 'Title with 🌟🎉'],
        'content': [
            'Content with 特殊字符 and symbols',
            'Content with ñáéíóú and accents',
            'Content with 🌟🎉 emojis'
        ]
    })
    df.to_excel(TEST_DATA_DIR / 'special_chars.xlsx', index=False)
    
    # CSV with special characters
    df.to_csv(TEST_DATA_DIR / 'special_chars.csv', index=False)
    
    # Word with special characters
    doc = Document()
    doc.add_heading('Document with 特殊字符', 0)
    doc.add_paragraph('Content with ñáéíóú')
    doc.add_paragraph('Content with 🌟🎉')
    doc.save(TEST_DATA_DIR / 'special_chars.docx')
    
    # PDF with special characters
    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    c.setTitle("PDF with 特殊字符")
    c.drawString(100, 750, "Content with ñáéíóú")
    c.drawString(100, 700, "Content with symbols")
    c.save()
    with open(TEST_DATA_DIR / 'special_chars.pdf', 'wb') as f:
        f.write(buffer.getvalue())

def create_mixed_encoding_files():
    """Create files with mixed encodings."""
    # CSV with mixed encodings
    content = [
        ['title', 'content'],
        ['UTF-8 Content', 'This is UTF-8 text with 特殊字符'],
        ['Latin1 Content', 'This is Latin1 text with ñáéíóú']
    ]
    
    # Write UTF-8 version
    with open(TEST_DATA_DIR / 'mixed_utf8.csv', 'w', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerows(content)
    
    # Write Latin1 version
    try:
        with open(TEST_DATA_DIR / 'mixed_latin1.csv', 'w', encoding='latin1') as f:
            writer = csv.writer(f)
            writer.writerows([
                ['title', 'content'],
                ['Latin1 Content', 'This is Latin1 text with áéíóú']
            ])
    except UnicodeEncodeError:
        # Fall back to ASCII if Latin1 encoding fails
        with open(TEST_DATA_DIR / 'mixed_latin1.csv', 'w', encoding='ascii', errors='ignore') as f:
            writer = csv.writer(f)
            writer.writerows([
                ['title', 'content'],
                ['ASCII Content', 'This is ASCII text only']
            ])

@pytest.mark.asyncio
async def test_empty_files():
    """Test processing of empty files."""
    factory = DocumentProcessorFactory()
    
    # Test empty Excel
    processor = factory.get_processor(TEST_DATA_DIR / 'empty.xlsx')
    results = processor.extract_content(TEST_DATA_DIR / 'empty.xlsx')
    assert len(results) == 0
    
    # Test empty CSV
    processor = factory.get_processor(TEST_DATA_DIR / 'empty.csv')
    results = processor.extract_content(TEST_DATA_DIR / 'empty.csv')
    assert len(results) == 0
    
    # Test empty Word
    processor = factory.get_processor(TEST_DATA_DIR / 'empty.docx')
    results = processor.extract_content(TEST_DATA_DIR / 'empty.docx')
    assert len(results) == 0
    
    # Test empty PDF
    processor = factory.get_processor(TEST_DATA_DIR / 'empty.pdf')
    results = processor.extract_content(TEST_DATA_DIR / 'empty.pdf')
    assert len(results) == 0

@pytest.mark.asyncio
async def test_malformed_files():
    """Test processing of malformed files."""
    factory = DocumentProcessorFactory()
    
    # Test malformed Excel
    processor = factory.get_processor(TEST_DATA_DIR / 'malformed.xlsx')
    results = processor.extract_content(TEST_DATA_DIR / 'malformed.xlsx')
    assert len(results) == 0
    
    # Test malformed CSV
    processor = factory.get_processor(TEST_DATA_DIR / 'malformed.csv')
    results = processor.extract_content(TEST_DATA_DIR / 'malformed.csv')
    assert len(results) == 0
    
    # Test malformed Word
    processor = factory.get_processor(TEST_DATA_DIR / 'malformed.docx')
    results = processor.extract_content(TEST_DATA_DIR / 'malformed.docx')
    assert len(results) == 0
    
    # Test malformed PDF
    processor = factory.get_processor(TEST_DATA_DIR / 'malformed.pdf')
    results = processor.extract_content(TEST_DATA_DIR / 'malformed.pdf')
    assert len(results) == 0

@pytest.mark.asyncio
async def test_large_files():
    """Test processing of large files."""
    factory = DocumentProcessorFactory()
    
    # Test large Excel
    processor = factory.get_processor(TEST_DATA_DIR / 'large.xlsx')
    results = processor.extract_content(TEST_DATA_DIR / 'large.xlsx')
    assert len(results) == 10000
    assert all('Title' in r['title'] for r in results)
    assert all('Content' in r['content'] for r in results)
    
    # Test large CSV
    processor = factory.get_processor(TEST_DATA_DIR / 'large.csv')
    results = processor.extract_content(TEST_DATA_DIR / 'large.csv')
    assert len(results) == 10000
    assert all('Title' in r['title'] for r in results)
    assert all('Content' in r['content'] for r in results)
    
    # Test large Word
    processor = factory.get_processor(TEST_DATA_DIR / 'large.docx')
    results = processor.extract_content(TEST_DATA_DIR / 'large.docx')
    assert len(results) == 1
    assert 'Large Document' in results[0]['title']
    assert 'Paragraph' in results[0]['content']
    
    # Test large PDF
    processor = factory.get_processor(TEST_DATA_DIR / 'large.pdf')
    results = processor.extract_content(TEST_DATA_DIR / 'large.pdf')
    assert len(results) == 1
    assert 'Page' in results[0]['content']

@pytest.mark.asyncio
async def test_special_characters():
    """Test processing of files with special characters."""
    factory = DocumentProcessorFactory()
    
    # Test Excel with special characters
    processor = factory.get_processor(TEST_DATA_DIR / 'special_chars.xlsx')
    results = processor.extract_content(TEST_DATA_DIR / 'special_chars.xlsx')
    assert len(results) == 3
    assert any('特殊字符' in r['title'] for r in results)
    assert any('ñáéíóú' in r['content'] for r in results)
    
    # Test CSV with special characters
    processor = factory.get_processor(TEST_DATA_DIR / 'special_chars.csv')
    results = processor.extract_content(TEST_DATA_DIR / 'special_chars.csv')
    assert len(results) == 3
    assert any('特殊字符' in r['title'] for r in results)
    assert any('ñáéíóú' in r['content'] for r in results)
    
    # Test Word with special characters
    processor = factory.get_processor(TEST_DATA_DIR / 'special_chars.docx')
    results = processor.extract_content(TEST_DATA_DIR / 'special_chars.docx')
    assert len(results) == 1
    assert '特殊字符' in results[0]['title']
    assert 'ñáéíóú' in results[0]['content']
    
    # Test PDF with special characters
    processor = factory.get_processor(TEST_DATA_DIR / 'special_chars.pdf')
    results = processor.extract_content(TEST_DATA_DIR / 'special_chars.pdf')
    assert len(results) == 1
    assert '特殊字符' in results[0]['title']
    assert 'ñáéíóú' in results[0]['content'].lower()

@pytest.mark.asyncio
async def test_mixed_encodings():
    """Test processing of files with mixed encodings."""
    processor = CSVProcessor()
    
    # Test UTF-8 CSV
    results = processor.extract_content(TEST_DATA_DIR / 'mixed_utf8.csv')
    assert len(results) == 2
    assert any('特殊字符' in r['content'] for r in results)
    
    # Test Latin1 CSV
    results = processor.extract_content(TEST_DATA_DIR / 'mixed_latin1.csv')
    assert len(results) == 1
    # Check if either Latin1 content or ASCII fallback is present
    assert any('áéíóú' in r['content'] or 'ASCII' in r['content'] for r in results)

def teardown_module(module):
    """Clean up test files after running tests."""
    for file in TEST_DATA_DIR.glob('*.*'):
        file.unlink()
    TEST_DATA_DIR.rmdir() 