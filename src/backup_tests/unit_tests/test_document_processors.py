"""
Tests for document processors.
"""

import os
import sys
import pytest
import pandas as pd
from pathlib import Path
from docx import Document
from openpyxl import Workbook
import pypdf
from reportlab.pdfgen import canvas
from io import BytesIO


# Ensure the src directory is in the path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Ensure the src directory is in the path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '.')))
from src.utils.document_processors import (
    DocumentProcessorFactory,
    ExcelProcessor,
    WordProcessor,
    PDFProcessor,
    CSVProcessor
)

# Test data directory
TEST_DATA_DIR = Path("test_data/documents")

def setup_module(module):
    """Set up test data before running tests."""
    TEST_DATA_DIR.mkdir(parents=True, exist_ok=True)
    
    # Create test Excel file
    create_test_excel()
    
    # Create test CSV file
    create_test_csv()
    
    # Create test Word file
    create_test_word()
    
    # Create test PDF file
    create_test_pdf()

def create_test_excel():
    """Create a test Excel file."""
    df = pd.DataFrame({
        'title': ['Article 1', 'Article 2'],
        'content': [
            'This is the content of article 1. It contains some facts.',
            'This is the content of article 2. It has different facts.'
        ]
    })
    df.to_excel(TEST_DATA_DIR / 'test.xlsx', index=False)

def create_test_csv():
    """Create a test CSV file."""
    df = pd.DataFrame({
        'heading': ['Document 1', 'Document 2'],
        'text': [
            'This is the text content of document 1.',
            'This is the text content of document 2.'
        ]
    })
    df.to_csv(TEST_DATA_DIR / 'test.csv', index=False)

def create_test_word():
    """Create a test Word document."""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph in the Word document.')
    doc.add_paragraph('It contains multiple paragraphs with content.')
    doc.save(TEST_DATA_DIR / 'test.docx')

def create_test_pdf():
    """Create a test PDF document."""
    buffer = BytesIO()
    c = canvas.Canvas(buffer)
    c.setTitle("Test PDF")
    c.drawString(100, 750, "This is a test PDF document.")
    c.drawString(100, 700, "It contains some test content for extraction.")
    c.save()
    
    # Write the PDF to file
    with open(TEST_DATA_DIR / 'test.pdf', 'wb') as f:
        f.write(buffer.getvalue())

@pytest.mark.asyncio
async def test_excel_processor():
    """Test Excel document processor."""
    processor = ExcelProcessor()
    file_path = TEST_DATA_DIR / 'test.xlsx'
    
    # Test file type detection
    assert processor.can_process(file_path)
    
    # Test content extraction
    results = processor.extract_content(file_path)
    assert len(results) == 2
    assert results[0]['title'] == 'Article 1'
    assert 'content of article 1' in results[0]['content']
    assert results[1]['title'] == 'Article 2'
    assert 'content of article 2' in results[1]['content']

@pytest.mark.asyncio
async def test_csv_processor():
    """Test CSV document processor."""
    processor = CSVProcessor()
    file_path = TEST_DATA_DIR / 'test.csv'
    
    # Test file type detection
    assert processor.can_process(file_path)
    
    # Test content extraction
    results = processor.extract_content(file_path)
    assert len(results) == 2
    assert results[0]['title'] == 'Document 1'
    assert 'content of document 1' in results[0]['content']
    assert results[1]['title'] == 'Document 2'
    assert 'content of document 2' in results[1]['content']

@pytest.mark.asyncio
async def test_word_processor():
    """Test Word document processor."""
    processor = WordProcessor()
    file_path = TEST_DATA_DIR / 'test.docx'
    
    # Test file type detection
    assert processor.can_process(file_path)
    
    # Test content extraction
    results = processor.extract_content(file_path)
    assert len(results) == 1
    assert results[0]['title'] == 'Test Document'
    assert 'test paragraph' in results[0]['content']
    assert 'multiple paragraphs' in results[0]['content']

@pytest.mark.asyncio
async def test_pdf_processor():
    """Test PDF document processor."""
    processor = PDFProcessor()
    file_path = TEST_DATA_DIR / 'test.pdf'
    
    # Test file type detection
    assert processor.can_process(file_path)
    
    # Test content extraction
    results = processor.extract_content(file_path)
    assert len(results) == 1
    assert results[0]['title'] == 'Test PDF'
    
    # Test content with normalized text
    content = results[0]['content'].lower()
    assert 'this is a test pdf document' in content
    assert 'test content for extraction' in content

@pytest.mark.asyncio
async def test_processor_factory():
    """Test document processor factory."""
    factory = DocumentProcessorFactory()
    
    # Test Excel processor selection
    excel_processor = factory.get_processor(TEST_DATA_DIR / 'test.xlsx')
    assert isinstance(excel_processor, ExcelProcessor)
    
    # Test CSV processor selection
    csv_processor = factory.get_processor(TEST_DATA_DIR / 'test.csv')
    assert isinstance(csv_processor, CSVProcessor)
    
    # Test Word processor selection
    word_processor = factory.get_processor(TEST_DATA_DIR / 'test.docx')
    assert isinstance(word_processor, WordProcessor)
    
    # Test PDF processor selection
    pdf_processor = factory.get_processor(TEST_DATA_DIR / 'test.pdf')
    assert isinstance(pdf_processor, PDFProcessor)
    
    # Test unsupported file type
    unsupported = factory.get_processor(TEST_DATA_DIR / 'test.txt')
    assert unsupported is None

def teardown_module(module):
    """Clean up test files after running tests."""
    for file in TEST_DATA_DIR.glob('test.*'):
        file.unlink()
    TEST_DATA_DIR.rmdir() 